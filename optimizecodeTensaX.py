import io
import os
import re
import sys
from langchain_groq import ChatGroq
import zipfile
import logging
import warnings
import gradio as gr
from PIL import Image
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader, PdfWriter
import matplotlib.pyplot as plt
from reportlab.lib import colors
from rich.console import Console
from rich.markdown import Markdown
from reportlab.pdfgen import canvas
from langchain_groq import ChatGroq
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import A4
from openai import OpenAI
from reportlab.platypus import Image as ReportLabImage, SimpleDocTemplate, Paragraph, Spacer, Frame, Table, TableStyle
from dotenv import load_dotenv
load_dotenv()



logging.getLogger().setLevel(logging.ERROR)
warnings.filterwarnings("ignore")


os.environ["OPENAI_API_KEY"]=os.getenv("OPENAI_API_KEY")
openai_api_key=os.environ["OPENAI_API_KEY"]
if not  openai_api_key:
    raise ValueError("Key not found")

client = OpenAI(
    api_key=openai_api_key
)


os.environ["GROQ_API_KEY"]=os.getenv("GROQ_API_KEY")
groq_api_key=os.environ["GROQ_API_KEY"]
if not  groq_api_key:
    raise ValueError("Key not found")


chat_groq = ChatGroq(model_name="llama3-groq-70b-8192-tool-use-preview", api_key=groq_api_key)

zip_file_path = '/Users/tusharsharma/Documents/TensaX Optimization Project/OneDrive_1_9-10-2024 (1).zip'
extract_dir = '/Users/tusharsharma/Documents/TensaX Optimization Project/images'
output_folder = '/Users/tusharsharma/Documents/TensaX Optimization Project/outputs'

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

def read_file(file_name, file_data):
    if file_data is None:
        return "No file uploaded."

    if file_name.endswith('.docx'):
        extracted_text = extract_text_from_docx(file_data)
        return extracted_text

    elif file_name.endswith('.pdf'):
        reader = PdfReader(io.BytesIO(file_data))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    else:
        return "Unsupported file type. Please upload a .docx or .pdf file."


def save_plot_as_image(output_folder, image_name):
    plt.savefig(os.path.join(output_folder, image_name), bbox_inches='tight', pad_inches=0.1)
    plt.close()

def display_image(image_name, folder_path):
    normalized_image_name = normalize_image_name(image_name)
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if normalize_image_name(file) == normalized_image_name:
                image_path = os.path.join(root, file)
                img = Image.open(image_path)
                img = img.convert("RGBA")
                white_bg = Image.new("RGBA", img.size, (255, 255, 255))
                img = Image.alpha_composite(white_bg, img).convert("RGB")
                save_plot_as_image(output_folder, normalized_image_name + '.png')
                return image_path
    return None

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    return '\n'.join(full_text)

def extract_images_from_zip(zip_file_path, extract_dir):
    if not os.path.exists(extract_dir):
        os.makedirs(extract_dir)
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)

def normalize_image_name(image_name):
    return image_name.lower().replace(' ', '_')

def extract_image_names_from_dir(dir_path):
    image_names = set()
    for root, dirs, files in os.walk(dir_path):
        for file in files:
            if file.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
                normalized_name = normalize_image_name(file)
                image_names.add(normalized_name)
    return image_names

def generate_llm_content_groq(text, image_tags):
    image_tags_list = list(image_tags)

    # Instead of splitting into single sentences, divide into logical paragraphs or sections
    paragraphs = text.strip().split("\n\n")  # Split by double newline for paragraphs

    all_paragraphs_prompt = ""
    for i, paragraph in enumerate(paragraphs):
        all_paragraphs_prompt += f"""
### Paragraph {i+1}
Content: "{paragraph}"
Simplify this paragraph into a few aphasia-friendly bullet points, ensuring that the key information and depth of meaning are preserved. Follow the same guidelines and format as before (single sentences, simple language, active voice, bold keywords, relevant images, '|’ separator).
"""

    prompt = f"""You are tasked with strictly creating an aphasia-friendly consent form by processing paragraphs of text.
### Objective:
Simplify the text specifically for individuals with aphasia, ensuring maximum ease of understanding while preserving the depth of information.
### Strict Guidelines (Follow each one carefully):
1. *Sentence Structure*: Each bullet point must be a single sentence, but it can be slightly longer (up to 8 words) to accommodate important details. Avoid overly long or complex sentences.
2. *Language*: Use only simple and clear words. No complex terms or jargon.
3. *Active Voice*: Use active voice exclusively. Limit or avoid pronouns.
4. *Keywords*: Each sentence **must* contain one bold keyword, formatted as *keyword*.
5. *Image Suggestions*: For each simplified point, always assign the **most relevant image* from the list: {', '.join(image_tags_list)}.
  - Each point *must* have an image, and it must be drawn from the provided image tags only. **Do not leave any point without an assigned image**.
6. *Formatting*: Separate each simplified point and its suggested image with a '|' character.
### Important:
- *No deviation from these guidelines is allowed*.
- *Revise your response* to meet these criteria if any point or rule is not followed.
- **Ensure that the simplified points capture the essential meaning and depth of the original text**.
### Example Format:
PERSONAL DETAILS
- Write your *name*. | name_tag.png
- Write your *birthdate*. | calendar.jpg
- Write your *hospital number*. | hospital_id.png
TREATMENT
- This treatment helps your arm *move*. | arm_movement.png
- You will do special *exercises*. | physiotherapy_exercises.jpg
- You might have some *pain* after. | pain_relief.png
IMPORTANT
- There are some *risks*. | warning_sign.jpg
- This treatment has many *benefits*. | thumbs_up.png
- You can choose other *treatments*. | doctor_consultation.png
- You need to give your *consent*. | signature.png
{all_paragraphs_prompt}  # Append all the paragraphs to the main prompt
"""

    response = chat_groq.invoke(prompt, temperature=0.0)

    if response:
        content = response.content.strip()
        if "Error" in content or not any(tag in content for tag in image_tags_list):
            return "Error: LLM did not return valid content or failed to follow guidelines. Please try again."
        return content
    else:
        return "Error: LLM did not return valid content."
    
    
def generate_llm_content_openai(text, image_tags):
        image_tags_list = list(image_tags)

        # Instead of splitting into single sentences, divide into logical paragraphs or sections
        paragraphs = text.strip().split("\n\n")  # Split by double newline for paragraphs

        all_paragraphs_prompt = ""
        for i, paragraph in enumerate(paragraphs):
            all_paragraphs_prompt += f"""
    ### Paragraph {i+1}
    Content: "{paragraph}"
    Simplify this paragraph into a few aphasia-friendly bullet points, ensuring that the key information and depth of meaning are preserved. Follow the same guidelines and format as before (single sentences, simple language, active voice, bold keywords, relevant images, '|’ separator).
    """

        prompt = f"""You are tasked with strictly creating an aphasia-friendly consent form by processing paragraphs of text.
    ### Objective:
    Simplify the text specifically for individuals with aphasia, ensuring maximum ease of understanding while preserving the depth of information.
    ### Strict Guidelines (Follow each one carefully):
    1. *Sentence Structure*: Each bullet point must be a single sentence, but it can be slightly longer (up to 8 words) to accommodate important details. Avoid overly long or complex sentences.
    2. *Language*: Use only simple and clear words. No complex terms or jargon.
    3. *Active Voice*: Use active voice exclusively. Limit or avoid pronouns.
    4. *Keywords*: Each sentence **must* contain one bold keyword, formatted as *keyword*.
    5. *Image Suggestions*: For each simplified point, always assign the **most relevant image* from the list: {', '.join(image_tags_list)}.
    - Each point *must* have an image, and it must be drawn from the provided image tags only. **Do not leave any point without an assigned image**.
    6. *Formatting*: Separate each simplified point and its suggested image with a '|' character.
    ### Important:
    - *No deviation from these guidelines is allowed*.
    - *Revise your response* to meet these criteria if any point or rule is not followed.
    - **Ensure that the simplified points capture the essential meaning and depth of the original text**.
    ### Example Format:
    PERSONAL DETAILS
    - Write your *name*. | name_tag.png
    - Write your *birthdate*. | calendar.jpg
    - Write your *hospital number*. | hospital_id.png
    TREATMENT
    - This treatment helps your arm *move*. | arm_movement.png
    - You will do special *exercises*. | physiotherapy_exercises.jpg
    - You might have some *pain* after. | pain_relief.png
    IMPORTANT
    - There are some *risks*. | warning_sign.jpg
    - This treatment has many *benefits*. | thumbs_up.png
    - You can choose other *treatments*. | doctor_consultation.png
    - You need to give your *consent*. | signature.png
    {all_paragraphs_prompt}  # Append all the paragraphs to the main prompt
    """

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}],
            temperature=0.0
        )



def display_consent_form(llm_content, folder_path):
    content_points = llm_content.split('\n')
    consent_data = []
    for item in content_points:
        simplified_text = item.strip()
        if not simplified_text:
            continue
        parts = simplified_text.split('|')
        consent_point = parts[0].strip()
        image_tag = parts[1].strip() if len(parts) > 1 else None
        image_path = None
        if image_tag:
            image_path = display_image(image_tag, folder_path)
        consent_data.append((consent_point, image_path))
    return consent_data

def save_as_docx(consent_data, output_file_path):
    doc = Document()
    count = 0
    for consent_point, image_path in consent_data:
        if count % 3 == 0:
            doc.add_page_break()

        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', consent_point)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(14)
            else:
                run = p.add_run(part)
                run.font.size = Pt(12)

        if image_path:
            doc.add_picture(image_path, width=Pt(300))

        count += 1

    doc.save(output_file_path)



def save_as_pdf(consent_data, output_file_path):
    pdf_writer = PdfWriter()

    for i in range(0, len(consent_data), 3):
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=A4)
        can.setFont("Helvetica", 12)

        y_position = 750

        for j in range(3):
            if i + j < len(consent_data):
                consent_point, image_path = consent_data[i + j]
                parts = re.split(r'(\*\*.*?\*\*)', consent_point)

                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        can.setFont("Helvetica-Bold", 14)
                        can.drawString(50, y_position, part[2:-2])
                    else:
                        can.setFont("Helvetica", 12)
                        can.drawString(50, y_position, part)
                    y_position -= 20

                if image_path:
                    can.drawImage(image_path, 50, y_position - 100, width=100, height=100)
                    y_position -= 120

                y_position -= 40

        can.save()
        packet.seek(0)
        new_pdf = PdfReader(packet)
        page = new_pdf.pages[0]
        pdf_writer.add_page(page)

    with open(output_file_path, 'wb') as output_file:
        pdf_writer.write(output_file)


import gradio as gr

def process_file(file, llm_choice):
    extract_images_from_zip(zip_file_path, extract_dir)
    consent_form_text = read_file(file.name, file)

    if "No file uploaded." in consent_form_text or "Unsupported file type." in consent_form_text:
        return consent_form_text

    image_names = extract_image_names_from_dir(extract_dir)

    # Call the corresponding LLM function based on the user's selection
    if llm_choice == "OpenAI":
        llm_content = generate_llm_content_openai(consent_form_text, image_names)
    elif llm_choice == "Groq":
        llm_content = generate_llm_content_groq(consent_form_text, image_names)

    consent_data = display_consent_form(llm_content, extract_dir)

    docx_path = os.path.join(output_folder, 'consent_form.docx')
    pdf_path = os.path.join(output_folder, 'consent_form.pdf')
    save_as_docx(consent_data, docx_path)
    save_as_pdf(consent_data, pdf_path)

    return docx_path, pdf_path

def download_file(file_name):
    return gr.File(file_name=file_name)

# Define the Gradio interface
with gr.Blocks() as demo:
    gr.Markdown("<h1 style='text-align: center;'>Developed by Aman Sahoo and Innovation Lab</h1>")

    # Add radio buttons for LLM model selection
    llm_choice = gr.Radio(choices=["OpenAI", "Groq"], label="Select LLM Model")

    gr.Interface(
        fn=process_file,
        inputs=[gr.File(label="Upload File"), llm_choice],  # Pass the radio button selection as input
        outputs=[gr.File(label="Download DOCX"), gr.File(label="Download PDF")]
    )

demo.launch(debug=True)



